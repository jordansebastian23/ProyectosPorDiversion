import os
from docx import Document
from docx.shared import Cm

carpeta = r'C:\Users\cpdsfpe\Music\Finiquitos\documentos-2025-05-08-10_37_32'
ruta_imagen = r'C:\Users\cpdsfpe\Music\Finiquitos\Imagen.png'

for archivo in os.listdir(carpeta):
    if archivo.endswith('.docx'):
        try:
            ruta_doc = os.path.join(carpeta, archivo)
            doc = Document(ruta_doc)

            seccion = doc.sections[0]
            encabezado = seccion.header

            for p in encabezado.paragraphs:
                p.clear()

            parrafo = encabezado.paragraphs[0]
            parrafo.alignment = 0
            run = parrafo.add_run()

            # Establece el tamaño de la imagen
            run.add_picture(ruta_imagen, width=Cm(1.56), height=Cm(1.86))

            doc.save(ruta_doc)
            print(f"✅ Imagen añadida a: {archivo}")
        except Exception as e:
            print(f"⚠️ Error en {archivo}: {e}")

print("Proceso completado con imágenes redimensionadas 📏✅")
